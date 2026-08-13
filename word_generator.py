from docx import Document
import os

def generate_word(addresses, filename="report.docx"):
    """
    Создает Word-документ со списком адресов.
    """
    doc = Document()
    doc.add_heading('Список адресов', 0)
    
    if not addresses:
        doc.add_paragraph('Адреса не найдены.')
    else:
        for i, addr in enumerate(addresses, 1):
            doc.add_paragraph(f'{i}. {addr}', style='List Bullet')
    
    # Сохраняем во временную папку exports
    output_path = os.path.join('exports', filename)
    doc.save(output_path)
    return output_path
